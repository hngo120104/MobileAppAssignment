from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


DOCX = Path(__file__).resolve().parents[1] / "BAO_CAO_APP_LOP_HOC_LAP_TRINH_NHI_HOAN_CHINH.docx"
doc = Document(DOCX)

paragraphs = [p for p in doc.paragraphs if p.text.strip()]
text = "\n".join(p.text for p in paragraphs)
h1 = [p.text for p in paragraphs if p.style.name == "Heading 1"]
h2 = [p.text for p in paragraphs if p.style.name == "Heading 2"]
page_breaks = 0
numbered_paragraphs = 0
manual_bullets = []

for paragraph in doc.paragraphs:
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            page_breaks += 1
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        numbered_paragraphs += 1
    if paragraph.text.strip().startswith(("- ", "• ")):
        manual_bullets.append(paragraph.text[:80])

required_phrases = [
    "MỤC LỤC",
    "DANH SÁCH TỪ VIẾT TẮT",
    "DANH SÁCH HÌNH VẼ",
    "DANH SÁCH BẢNG BIỂU",
    "Kiến trúc tổng quan",
    "Biểu đồ use case tổng quan",
    "Use case chi tiết",
    "Biểu đồ lớp",
    "Biểu đồ tuần tự",
    "Sơ đồ thực thể quan hệ",
    "Mô hình triển khai",
    "Các bước cài đặt và triển khai",
    "Kết quả chức năng",
    "Kết quả kiểm thử",
    "KẾT LUẬN, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN",
    "TÀI LIỆU THAM KHẢO",
]
missing = [phrase for phrase in required_phrases if phrase not in text]

with ZipFile(DOCX) as archive:
    bad_member = archive.testzip()
    numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
    document_xml = archive.read("word/document.xml").decode("utf-8")

assert bad_member is None, f"Corrupt member: {bad_member}"
assert not missing, f"Missing required sections: {missing}"
assert len(h1) == 5, f"Unexpected Heading 1 count: {len(h1)}"
assert len(h2) >= 30, f"Too few Heading 2 paragraphs: {len(h2)}"
assert page_breaks == 23, f"Expected 23 explicit page breaks, got {page_breaks}"
assert numbered_paragraphs >= 40, "Real numbered/bulleted paragraphs are missing"
assert not manual_bullets, f"Fake bullets found: {manual_bullets}"
assert numbering_xml.count("w:abstractNum") >= 2, "Custom numbering definitions missing"
assert document_xml.count("classDiagram") == 1
assert document_xml.count("sequenceDiagram") == 3
assert document_xml.count("erDiagram") == 1
assert document_xml.count("flowchart") >= 3
assert len(doc.tables) == 8

normal = doc.styles["Normal"].paragraph_format
h1_style = doc.styles["Heading 1"].paragraph_format
assert round(normal.space_after.pt, 1) == 6.0
assert doc.styles["Normal"].font.name == "Times New Roman"
assert round(doc.styles["Normal"].font.size.pt, 1) == 13.0
assert round(float(normal.line_spacing), 2) == 1.30
assert round(h1_style.space_before.pt, 1) == 16.0
assert round(h1_style.space_after.pt, 1) == 10.0

captions = [p.text for p in paragraphs if p.style.name == "Caption"]
assert sum(text.startswith("Hình ") for text in captions) == 8
assert sum(text.startswith("Bảng ") for text in captions) == 6
for chapter, expected in {"1": 2, "2": 9, "3": 3}.items():
    chapter_captions = [text for text in captions if text.startswith((f"Hình {chapter}.", f"Bảng {chapter}."))]
    assert len(chapter_captions) == expected, (chapter, chapter_captions)

for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.trPr
        if tr_pr is not None:
            heights = tr_pr.findall(qn("w:trHeight"))
            for height in heights:
                assert height.get(qn("w:hRule")) != "exact", "Fixed row height found"

print(f"DOCX: {DOCX}")
print(f"Non-empty paragraphs: {len(paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print(f"Heading 1 / Heading 2: {len(h1)} / {len(h2)}")
print(f"Explicit page breaks: {page_breaks} (minimum 24 pages by authored structure)")
print(f"Real numbered/list paragraphs: {numbered_paragraphs}")
print("Required content, ZIP integrity, styles, Mermaid blocks, numbering and row-height audit: PASS")
