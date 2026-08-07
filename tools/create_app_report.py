from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "BAO_CAO_APP_LOP_HOC_LAP_TRINH_NHI_HOAN_CHINH.docx"

FONT_NAME = "Times New Roman"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
GRAY = "5B6573"
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F7F8FA"
CALLOUT_FILL = "EEF4FA"
WHITE = "FFFFFF"
BLACK = "000000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name=FONT_NAME, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color=BLACK, bold=False, italic=False):
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.font.italic = italic


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    values = {
        "top": CELL_MARGIN_TOP_BOTTOM,
        "bottom": CELL_MARGIN_TOP_BOTTOM,
        "start": CELL_MARGIN_START_END,
        "end": CELL_MARGIN_START_END,
    }
    for edge, value in values.items():
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
    set_cell_margins(table)


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_NAME)
    fonts.set(qn("w:hAnsi"), FONT_NAME)
    fonts.set(qn("w:eastAsia"), FONT_NAME)
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "26")
    r_pr.append(size)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=11, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, display, end])


def configure_document(doc):
    section = doc.sections[0]
    # Named academic-format override: A4 with a wider binding margin.
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.59)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, FONT_NAME, 13)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.30
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = doc.styles["Title"]
    set_style_font(title, FONT_NAME, 24, NAVY, True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, FONT_NAME, 15, GRAY)
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading_tokens = {
        "Heading 1": (18, NAVY, 16, 10),
        "Heading 2": (15, BLUE, 12, 6),
        "Heading 3": (13, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        set_style_font(style, FONT_NAME, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = doc.styles["Caption"]
    set_style_font(caption, FONT_NAME, 13, GRAY, False, True)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("Mermaid Code", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, FONT_NAME, 10.5, NAVY)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.keep_together = True

    note = doc.styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(note, FONT_NAME, 13, DARK_BLUE, True)
    note.paragraph_format.left_indent = Inches(0.15)
    note.paragraph_format.right_indent = Inches(0.15)
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.20

    front_heading = doc.styles.add_style("Front Matter Heading", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(front_heading, FONT_NAME, 16, NAVY, True)
    front_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front_heading.paragraph_format.space_before = Pt(0)
    front_heading.paragraph_format.space_after = Pt(16)

    index_entry = doc.styles.add_style("Static Index Entry", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(index_entry, FONT_NAME, 13, BLACK)
    index_entry.paragraph_format.space_before = Pt(0)
    index_entry.paragraph_format.space_after = Pt(3)
    index_entry.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("BÁO CÁO ỨNG DỤNG ANDROID  |  LỚP HỌC LẬP TRÌNH NHÍ")
    set_run_font(run, size=10.5, color=GRAY, bold=True)
    add_page_number(section.footer.paragraphs[0])

    return add_numbering_definition(doc, "bullet"), add_numbering_definition(doc, "decimal")


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text, bullet_num_id):
    paragraph = doc.add_paragraph(style="Normal")
    apply_numbering(paragraph, bullet_num_id)
    paragraph.add_run(text)
    return paragraph


def add_numbered(doc, text, number_num_id):
    paragraph = doc.add_paragraph(style="Normal")
    apply_numbering(paragraph, number_num_id)
    paragraph.add_run(text)
    return paragraph


def add_callout(doc, label, text):
    paragraph = doc.add_paragraph(style="Lead Callout")
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shading)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=DARK_BLUE)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, bold=False, color=NAVY)
    return paragraph


def add_mermaid(doc, caption_text, source):
    doc.add_paragraph(caption_text, style="Caption")
    paragraph = doc.add_paragraph(style="Mermaid Code")
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shading)
    run = paragraph.add_run(source.strip())
    set_run_font(run, FONT_NAME, 10.5, NAVY)


def add_table(doc, headers, rows, widths, font_size=13):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    for row_values in rows:
        row = table.add_row()
        for index, text in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(text)
            set_run_font(run, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_captioned_table(doc, caption_text, headers, rows, widths):
    doc.add_paragraph(caption_text, style="Caption")
    return add_table(doc, headers, rows, widths, 13)


def start_front_page(doc, title):
    doc.add_page_break()
    doc.add_paragraph(title, style="Front Matter Heading")


def add_index_entry(doc, label, page, level=0):
    paragraph = doc.add_paragraph(style="Static Index Entry")
    paragraph.paragraph_format.left_indent = Inches(0.25 * level)
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        p_pr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9000")
    tabs.append(tab)
    paragraph.add_run(label)
    paragraph.add_run("\t")
    paragraph.add_run(str(page))
    return paragraph


def start_page(doc, title, level=1):
    doc.add_page_break()
    doc.add_heading(title, level=level)


def add_use_case(doc, code, name, actor, precondition, steps, alternatives, postcondition, number_num_id):
    doc.add_heading(f"{code} — {name}", level=2)
    add_body(doc, f"Tác nhân: {actor}", "Tác nhân:")
    add_body(doc, f"Tiền điều kiện: {precondition}", "Tiền điều kiện:")
    for step in steps:
        add_numbered(doc, step, number_num_id)
    add_body(doc, f"Ngoại lệ: {alternatives}", "Ngoại lệ:")
    add_body(doc, f"Hậu điều kiện: {postcondition}", "Hậu điều kiện:")


def build_report():
    doc = Document()
    bullet_id, number_id = configure_document(doc)
    properties = doc.core_properties
    properties.title = "Báo cáo ứng dụng Lớp học lập trình nhí"
    properties.subject = "Phân tích, thiết kế, triển khai và kiểm thử ứng dụng Android MVVM"
    properties.author = "[HỌ VÀ TÊN SINH VIÊN]"
    properties.keywords = "Android, Java, SQLite, MVVM, Clean Architecture"

    # Trang bìa
    doc.add_paragraph().paragraph_format.space_after = Pt(70)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("BÁO CÁO ĐỒ ÁN ỨNG DỤNG ANDROID")
    set_run_font(run, size=13, color=BLUE, bold=True)
    doc.add_paragraph("LỚP HỌC LẬP TRÌNH NHÍ", style="Title")
    doc.add_paragraph(
        "Ứng dụng quản lý học viên, khóa học, ghi danh và báo cáo",
        style="Subtitle",
    )
    add_table(
        doc,
        ["THÔNG TIN", "NỘI DUNG"],
        [
            ("Mã sinh viên", "[MÃ SINH VIÊN]"),
            ("Họ và tên", "[HỌ VÀ TÊN SINH VIÊN]"),
            ("Tên ứng dụng", "Lớp học lập trình nhí"),
            ("Nền tảng", "Android • Java • SQLite"),
            ("Kiến trúc", "MVVM kết hợp phân tầng Clean Architecture"),
        ],
        [2500, 6860],
        13,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Tháng 08 năm 2026"), size=13, color=GRAY, italic=True)

    # Phần đầu báo cáo
    start_front_page(doc, "MỤC LỤC")
    for label, page, level in [
        ("DANH SÁCH TỪ VIẾT TẮT", 3, 0),
        ("DANH SÁCH HÌNH VẼ", 4, 0),
        ("DANH SÁCH BẢNG BIỂU", 5, 0),
        ("TÓM TẮT", 6, 0),
        ("CHƯƠNG 1. GIỚI THIỆU VÀ PHÂN TÍCH YÊU CẦU", 7, 0),
        ("1.1. Bối cảnh và bài toán", 7, 1),
        ("1.2. Mục tiêu", 7, 1),
        ("1.3. Phạm vi", 7, 1),
        ("1.4. Phân tích yêu cầu chức năng", 8, 1),
        ("1.5. Yêu cầu phi chức năng", 8, 1),
        ("1.6. Tác nhân và quy tắc nghiệp vụ", 10, 1),
        ("CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 12, 0),
        ("2.1. Kiến trúc tổng quan", 12, 1),
        ("2.2. Trách nhiệm các tầng", 12, 1),
        ("2.3. Biểu đồ use case tổng quan", 13, 1),
        ("2.4. Use case chi tiết - quản lý danh mục", 14, 1),
        ("2.5. Use case chi tiết - ghi danh và báo cáo", 15, 1),
        ("2.6. Biểu đồ lớp", 17, 1),
        ("2.7. Biểu đồ tuần tự", 18, 1),
        ("2.8. Luồng thực hiện báo cáo", 19, 1),
        ("2.9. Sơ đồ thực thể quan hệ và thiết kế dữ liệu", 20, 1),
        ("CHƯƠNG 3. TRIỂN KHAI, KẾT QUẢ VÀ ĐÁNH GIÁ", 22, 0),
        ("3.1. Mô hình triển khai", 22, 1),
        ("3.2. Yêu cầu môi trường", 22, 1),
        ("3.3. Các bước cài đặt và triển khai", 23, 1),
        ("3.4. Kết quả chức năng", 24, 1),
        ("3.5. Kết quả kiểm thử và đánh giá", 26, 1),
        ("KẾT LUẬN, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN", 27, 0),
        ("TÀI LIỆU THAM KHẢO", 28, 0),
    ]:
        add_index_entry(doc, label, page, level)

    start_front_page(doc, "DANH SÁCH TỪ VIẾT TẮT")
    add_table(
        doc,
        ["TỪ VIẾT TẮT", "TIẾNG ANH / TÊN ĐẦY ĐỦ", "Ý NGHĨA"],
        [
            ("API", "Application Programming Interface", "Giao diện lập trình ứng dụng"),
            ("APK", "Android Package Kit", "Gói cài đặt ứng dụng Android"),
            ("CRUD", "Create - Read - Update - Delete", "Thêm - xem - sửa - xóa dữ liệu"),
            ("DAO", "Data Access Object", "Đối tượng truy cập dữ liệu"),
            ("DB", "Database", "Cơ sở dữ liệu"),
            ("DI", "Dependency Injection", "Tiêm phụ thuộc"),
            ("ER", "Entity Relationship", "Mô hình thực thể - quan hệ"),
            ("MVVM", "Model - View - ViewModel", "Kiến trúc trình bày của ứng dụng"),
            ("SQL", "Structured Query Language", "Ngôn ngữ truy vấn có cấu trúc"),
            ("UI", "User Interface", "Giao diện người dùng"),
            ("UML", "Unified Modeling Language", "Ngôn ngữ mô hình hóa thống nhất"),
        ],
        [1500, 3900, 3960],
    )

    start_front_page(doc, "DANH SÁCH HÌNH VẼ")
    for label, page in [
        ("Hình 2.1 - Kiến trúc tổng quan", 12),
        ("Hình 2.2 - Biểu đồ use case tổng quan", 13),
        ("Hình 2.3 - Biểu đồ lớp theo các tầng", 17),
        ("Hình 2.4 - Tuần tự thêm/sửa học viên", 18),
        ("Hình 2.5 - Tuần tự ghi danh nhiều học viên", 18),
        ("Hình 2.6 - Tuần tự thực hiện báo cáo", 19),
        ("Hình 2.7 - Sơ đồ thực thể quan hệ", 20),
        ("Hình 3.1 - Mô hình triển khai", 22),
    ]:
        add_index_entry(doc, label, page)

    start_front_page(doc, "DANH SÁCH BẢNG BIỂU")
    for label, page in [
        ("Bảng 1.1 - Yêu cầu chức năng", 8),
        ("Bảng 1.2 - Công nghệ sử dụng", 11),
        ("Bảng 2.1 - Trách nhiệm của các tầng kiến trúc", 12),
        ("Bảng 2.2 - Ràng buộc dữ liệu", 20),
        ("Bảng 3.1 - Kết quả thực hiện chức năng", 24),
        ("Bảng 3.2 - Kết quả kiểm thử", 26),
    ]:
        add_index_entry(doc, label, page)

    start_front_page(doc, "TÓM TẮT")
    add_body(doc, "Báo cáo trình bày quá trình phân tích, thiết kế và triển khai ứng dụng Android “Lớp học lập trình nhí”. Ứng dụng hỗ trợ quản lý học viên nhỏ tuổi, khóa học Scratch/Python, ghi danh nhiều học viên và tám hình thức báo cáo/truy vấn. Dữ liệu được lưu cục bộ bằng SQLite; không cần máy chủ hay kết nối Internet.")
    add_body(doc, "Về kỹ thuật, hệ thống áp dụng MVVM kết hợp phân tầng. Activity chỉ nhận thao tác và render; ViewModel giữ trạng thái; use case xử lý validation/nghiệp vụ; repository interface tách domain khỏi SQLite; data layer đảm nhiệm SQL, mapping và transaction. Các tác vụ I/O chạy trên worker thread, sau đó trả kết quả về main thread qua LiveData.")
    add_callout(doc, "Phạm vi", "Báo cáo tập trung vào bản Android chạy độc lập. Các sơ đồ được cung cấp dưới dạng Mermaid để có thể dựng lại bằng Visual Paradigm.")

    # Trang 3
    start_page(doc, "CHƯƠNG 1. GIỚI THIỆU VÀ PHÂN TÍCH YÊU CẦU")
    doc.add_heading("1.1. Bối cảnh và bài toán", level=2)
    add_body(doc, "Các lớp lập trình cho trẻ thường quản lý đồng thời thông tin học viên, khóa học theo ngôn ngữ/cấp độ và quan hệ ghi danh. Khi dữ liệu được lưu rời rạc, việc kiểm tra một học viên đang thuộc khóa nào, tổng số học viên của từng khóa hoặc lọc nhóm học viên theo tiêu chí trở nên tốn thời gian và dễ sai sót.")
    add_body(doc, "Ứng dụng được xây dựng nhằm cung cấp một công cụ gọn nhẹ chạy trực tiếp trên thiết bị Android. Người quản lý có thể thao tác không cần tài khoản và không phụ thuộc mạng. SQLite phù hợp với quy mô bài toán vì hỗ trợ constraint, khóa ngoại, index và transaction ngay trên thiết bị.")
    doc.add_heading("1.2. Mục tiêu", level=2)
    for text in [
        "Quản lý đầy đủ vòng đời học viên và khóa học.",
        "Ghi danh hoặc hủy ghi danh nhiều học viên theo cơ chế all-or-nothing.",
        "Cung cấp báo cáo theo khóa, tình trạng ghi danh, độ tuổi, ngôn ngữ, cấp độ và thời gian.",
        "Thiết kế mã nguồn dễ kiểm thử, bảo trì và mở rộng theo MVVM.",
        "Đảm bảo dữ liệu nhất quán bằng validation ở domain và constraint ở database.",
    ]:
        add_bullet(doc, text, bullet_id)
    doc.add_heading("1.3. Phạm vi", level=2)
    add_body(doc, "Phiên bản hiện tại là ứng dụng Android một người dùng, lưu dữ liệu cục bộ. Hệ thống chưa có đăng nhập, đồng bộ đám mây, phân quyền, API máy chủ hay xuất báo cáo thành tệp.")

    # Trang 4
    start_page(doc, "1.4. Phân tích yêu cầu chức năng", level=2)
    add_captioned_table(
        doc,
        "Bảng 1.1 - Yêu cầu chức năng",
        ["MÃ", "CHỨC NĂNG", "KẾT QUẢ MONG ĐỢI"],
        [
            ("FR-01", "Quản lý học viên", "Thêm, sửa, xóa một/nhiều học viên; kiểm tra mã, tên, tuổi 5–18."),
            ("FR-02", "Quản lý khóa học", "CRUD khóa Scratch/Python; lưu cấp độ, ngày bắt đầu và ngày kết thúc."),
            ("FR-03", "Ghi danh", "Chọn khóa, lọc học viên chưa thuộc khóa và ghi danh nhiều học viên."),
            ("FR-04", "Hủy ghi danh", "Xem học viên trong khóa và hủy một hoặc nhiều quan hệ ghi danh."),
            ("FR-05", "Báo cáo theo khóa", "Liệt kê học viên của khóa được chọn và hiển thị tổng số."),
            ("FR-06", "Tổng hợp khóa học", "Liệt kê mọi khóa và số học viên, kể cả khóa có số lượng bằng 0."),
            ("FR-07", "Tra cứu học viên", "Tìm học viên chưa ghi danh hoặc lọc theo khoảng tuổi tùy chọn."),
            ("FR-08", "Thống kê khóa học", "Thống kê số khóa và lượt ghi danh theo ngôn ngữ hoặc cấp độ."),
            ("FR-09", "Lọc theo thời gian", "Liệt kê ghi danh trong khoảng ngày do người dùng nhập."),
            ("FR-10", "Truy vấn nhanh", "Tìm học viên 10–12 tuổi đã ghi danh khóa Python cơ bản."),
        ],
        [1050, 2500, 5810],
    )
    doc.add_heading("1.5. Yêu cầu phi chức năng", level=2)
    for text in [
        "Hiệu năng: không thực hiện SQLite I/O trên main thread; danh sách không phát sinh N+1 query.",
        "Tin cậy: dùng foreign key, unique constraint và transaction để tránh dữ liệu mồ côi hoặc thành công một phần.",
        "Khả dụng: hỗ trợ loading, empty state, validation tại đúng trường nhập và thông báo lỗi một lần.",
        "Bảo trì: chiều phụ thuộc presentation → domain ← data; Activity không biết SQLite/repository.",
        "Tương thích: minSdk 24, targetSdk 36; mã nguồn Java 11.",
    ]:
        add_bullet(doc, text, bullet_id)

    # Trang 5
    start_page(doc, "1.6. Tác nhân và quy tắc nghiệp vụ", level=2)
    doc.add_heading("Tác nhân", level=2)
    add_body(doc, "Hệ thống có một tác nhân chính là Người quản lý lớp học. Tác nhân này vận hành toàn bộ chức năng trên thiết bị Android. SQLite là thành phần lưu trữ nội bộ, không phải tác nhân bên ngoài.")
    doc.add_heading("Quy tắc nghiệp vụ", level=2)
    for text in [
        "Mã học viên và mã khóa học là duy nhất.",
        "Tuổi học viên nằm trong khoảng từ 5 đến 18.",
        "Ngôn ngữ khóa học chỉ nhận Scratch hoặc Python.",
        "Ngày bắt đầu khóa học không được sau ngày kết thúc; định dạng ngày là yyyy-MM-dd.",
        "Một học viên chỉ được ghi danh một lần trong cùng khóa.",
        "Thời điểm ghi danh được hệ thống tự động ghi nhận khi tạo quan hệ.",
        "Xóa học viên/khóa học đồng thời xóa các ghi danh liên quan.",
        "Batch ghi danh, hủy ghi danh và xóa nhiều phải thành công toàn bộ hoặc rollback toàn bộ.",
    ]:
        add_bullet(doc, text, bullet_id)
    doc.add_page_break()
    doc.add_heading("Công nghệ lựa chọn", level=2)
    add_captioned_table(
        doc,
        "Bảng 1.2 - Công nghệ sử dụng",
        ["THÀNH PHẦN", "LỰA CHỌN", "LÝ DO"],
        [
            ("Ngôn ngữ", "Java 11", "Phù hợp yêu cầu môn học và Android SDK."),
            ("Giao diện", "Android XML + Activity", "Dễ kiểm soát layout và lifecycle cơ bản."),
            ("State", "ViewModel + LiveData", "Giữ trạng thái qua configuration change."),
            ("Lưu trữ", "SQLiteOpenHelper", "Cục bộ, transaction và constraint đầy đủ."),
            ("Kiểm thử", "JUnit + AndroidX Test", "Tách unit test domain và instrumented test database."),
        ],
        [1700, 2200, 5460],
    )

    # Trang 6
    start_page(doc, "CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")
    doc.add_heading("2.1. Kiến trúc tổng quan", level=2)
    add_body(doc, "Ứng dụng không có server/API. Đây là mô hình standalone client: toàn bộ UI, nghiệp vụ và persistence nằm trên thiết bị nhưng được phân tách bằng interface và use case. Composition root là nơi duy nhất ghép concrete SQLite repository vào domain.")
    add_mermaid(doc, "Hình 2.1 - Kiến trúc tổng quan", """
flowchart LR
    U[Người quản lý] --> V[Activity / View]
    V --> VM[ViewModel]
    VM --> UC[Use Case]
    UC --> RI[ClassroomRepository interface]
    DR[SqliteClassroomRepository] -.implements.-> RI
    DR --> DB[(SQLite)]
    APP[Application + Factory] -.inject.-> VM
    APP -.compose.-> DR
""")
    add_callout(doc, "Chiều phụ thuộc", "presentation → domain ← data; package di/composition root được phép biết các concrete type để lắp ghép runtime.")
    doc.add_heading("2.2. Trách nhiệm các tầng", level=2)
    add_captioned_table(
        doc,
        "Bảng 2.1 - Trách nhiệm của các tầng kiến trúc",
        ["TẦNG", "TRÁCH NHIỆM", "KHÔNG ĐƯỢC CHỨA"],
        [
            ("View", "Nhận input, gửi action, observe và render.", "SQL, repository, validator, worker thread."),
            ("ViewModel", "UI state/event, điều phối use case, threading.", "Android View, SQLite implementation."),
            ("Domain", "Model, validation, use case, repository contract.", "Context, Cursor, Activity."),
            ("Data", "SQL, mapping, transaction, migration.", "Thông báo UI và ViewModel."),
            ("DI", "Tạo repository, use case và ViewModel.", "Nghiệp vụ hoặc render UI."),
        ],
        [1550, 3900, 3910],
    )

    # Trang 7
    start_page(doc, "2.3. Biểu đồ use case tổng quan", level=2)
    add_mermaid(doc, "Hình 2.2 - Biểu đồ use case tổng quan", """
flowchart LR
    A([Người quản lý lớp học])
    A --> UC1((Quản lý học viên))
    A --> UC2((Quản lý khóa học))
    A --> UC3((Ghi danh học viên))
    A --> UC4((Hủy ghi danh))
    A --> UC5((Báo cáo theo khóa))
    A --> UC6((Tra cứu chưa ghi danh / độ tuổi))
    A --> UC7((Thống kê ngôn ngữ / cấp độ))
    A --> UC8((Lọc ghi danh theo thời gian))
    A --> UC9((Truy vấn Python cơ bản 10-12 tuổi))
    UC1 -.include.-> V1((Kiểm tra dữ liệu học viên))
    UC2 -.include.-> V2((Kiểm tra dữ liệu khóa học))
    UC3 -.include.-> C1((Lọc học viên khả dụng))
    UC3 -.include.-> TX((Transaction nhiều bản ghi))
    UC4 -.include.-> TX
""")
    doc.add_heading("Phân rã chức năng", level=2)
    for text in [
        "Quản lý học viên: xem danh sách; thêm; sửa; xóa một; chọn và xóa nhiều.",
        "Quản lý khóa học: xem danh sách; thêm; sửa; xóa một; chọn và xóa nhiều.",
        "Ghi danh: chọn khóa; lấy ứng viên; chọn nhiều; xác nhận; cập nhật số lượng.",
        "Hủy ghi danh: xem danh sách trong khóa; chọn nhiều; xác nhận; cập nhật lại.",
        "Báo cáo: chọn một trong tám loại, nhập bộ lọc phù hợp và xem kết quả tại ReportActivity riêng.",
    ]:
        add_bullet(doc, text, bullet_id)

    # Trang 8
    start_page(doc, "2.4. Use case chi tiết - quản lý danh mục", level=2)
    add_use_case(
        doc, "UC-01", "Quản lý học viên", "Người quản lý",
        "Ứng dụng đã khởi động và database sẵn sàng.",
        [
            "Mở màn hình Học viên và tải danh sách theo tên.",
            "Chọn thêm mới hoặc chọn một dòng để sửa.",
            "Nhập mã, họ tên, tuổi và trình độ Scratch rồi lưu.",
            "ViewModel yêu cầu StudentUseCase kiểm tra và tạo Student.",
            "Repository insert/update; ViewModel phát event và tải lại danh sách.",
        ],
        "Thiếu trường hoặc tuổi không hợp lệ: hiển thị lỗi đúng input. Mã trùng/lỗi lưu: giữ dialog và phát thông báo.",
        "Danh sách phản ánh dữ liệu mới; không tạo bản ghi trùng mã.", number_id,
    )
    add_use_case(
        doc, "UC-02", "Quản lý khóa học", "Người quản lý",
        "Ứng dụng đã khởi động.",
        [
            "Mở màn hình Khóa học.",
            "Thêm/sửa mã, tên, ngôn ngữ, cấp độ và thời gian khóa học.",
            "CourseUseCase kiểm tra dữ liệu và tạo Course.",
            "Repository lưu rồi tải lại danh sách.",
        ],
        "Mã/tên/ngày trống; ngày sai định dạng hoặc thứ tự; mã trùng; bản ghi đã bị xóa.",
        "Khóa học được lưu nhất quán; xóa khóa sẽ cascade ghi danh.", number_id,
    )

    # Trang 9
    start_page(doc, "2.5. Use case chi tiết - ghi danh và báo cáo", level=2)
    add_use_case(
        doc, "UC-03", "Ghi danh nhiều học viên", "Người quản lý",
        "Tồn tại ít nhất một khóa học và học viên chưa thuộc khóa.",
        [
            "Chọn chức năng Ghi danh và chọn khóa học.",
            "EnrollmentUseCase lọc học viên chưa thuộc khóa.",
            "Người dùng chọn một hoặc nhiều học viên và xác nhận.",
            "Repository mở transaction và insert toàn bộ quan hệ.",
            "Nếu thành công, commit và cập nhật CourseSummary.",
        ],
        "Không có học viên; tất cả đã ghi danh; không chọn học viên; constraint lỗi khiến rollback.",
        "Không tồn tại ghi danh trùng và batch không thành công một phần.", number_id,
    )
    add_use_case(
        doc, "UC-04", "Hủy ghi danh", "Người quản lý",
        "Khóa học có học viên đã ghi danh.",
        [
            "Chạm khóa để tải danh sách học viên.",
            "Chọn một hoặc nhiều học viên cần hủy.",
            "Xác nhận và thực hiện transaction delete.",
            "Tải lại danh sách và số lượng học viên.",
        ],
        "Không chọn học viên hoặc quan hệ không còn tồn tại: không thay đổi dữ liệu.",
        "Các quan hệ được chọn bị xóa toàn bộ hoặc rollback.", number_id,
    )
    add_use_case(
        doc, "UC-05", "Xem báo cáo và truy vấn", "Người quản lý",
        "Database đã sẵn sàng; với báo cáo có tham số, người dùng nhập bộ lọc hợp lệ.",
        [
            "Mở ReportActivity và chọn loại báo cáo.",
            "Nhập khóa học, khoảng tuổi hoặc khoảng ngày nếu loại báo cáo yêu cầu.",
            "ReportViewModel yêu cầu ReportUseCase kiểm tra và chuẩn hóa bộ lọc.",
            "Repository ủy quyền truy vấn cho EnrollmentDao hoặc ReportDao.",
            "ViewModel ánh xạ domain model thành ReportRow và cập nhật LiveData.",
        ],
        "Thiếu khóa học, khoảng tuổi ngoài 5-18, ngày sai định dạng hoặc ngày bắt đầu sau ngày kết thúc: hiển thị lỗi và không chạy SQL.",
        "Danh sách, số lượng hoặc thống kê được hiển thị; có empty state khi không có kết quả.", number_id,
    )

    # Trang 10
    start_page(doc, "2.6. Biểu đồ lớp", level=2)
    add_mermaid(doc, "Hình 2.3 - Biểu đồ lớp theo các tầng", """
classDiagram
    class BaseMvvmListActivity
    class StudentActivity
    class StudentViewModel
    class StudentUseCase
    class ClassroomRepository
    class SqliteClassroomRepository
    class ClassroomDatabase
    class Student
    class Course
    class CourseSummary
    class CourseStatistic
    class Enrollment
    class ReportDao
    StudentActivity --|> BaseMvvmListActivity
    StudentActivity --> StudentViewModel : action / observe
    StudentViewModel --> StudentUseCase
    StudentUseCase --> ClassroomRepository
    SqliteClassroomRepository ..|> ClassroomRepository
    SqliteClassroomRepository --> ClassroomDatabase
    StudentUseCase --> Student
    ClassroomRepository --> Student
    ClassroomRepository --> Course
    ClassroomRepository --> CourseSummary
    ClassroomRepository --> CourseStatistic
    ClassroomRepository --> Enrollment
    SqliteClassroomRepository --> ReportDao
""")
    add_body(doc, "CourseViewModel, EnrollmentViewModel và ReportViewModel áp dụng cùng mẫu với use case tương ứng. BaseListViewModel quản lý state danh sách, loading, selection, executor và error event; UiEvent bảo đảm side effect không phát lại sau rotation.")
    doc.add_heading("Quan hệ quan trọng", level=2)
    for text in [
        "Activity phụ thuộc ViewModel, không phụ thuộc repository.",
        "ViewModel nhận use case qua constructor do ClassroomViewModelFactory tạo.",
        "Use case phụ thuộc abstraction ClassroomRepository.",
        "SqliteClassroomRepository implement abstraction và là nơi duy nhất truy cập SQLite.",
    ]:
        add_bullet(doc, text, bullet_id)

    # Trang 11
    start_page(doc, "2.7. Biểu đồ tuần tự", level=2)
    add_mermaid(doc, "Hình 2.4 - Tuần tự thêm/sửa học viên", """
sequenceDiagram
    actor U as Người quản lý
    participant A as StudentActivity
    participant VM as StudentViewModel
    participant UC as StudentUseCase
    participant R as Repository
    participant DB as SQLite
    U->>A: Nhập form và chọn Lưu
    A->>VM: saveStudent(raw input)
    VM->>UC: prepareStudent(...)
    alt Dữ liệu không hợp lệ
        UC-->>VM: validation error
        VM-->>A: error code
        A->>A: setError(EditText)
    else Hợp lệ
        VM->>UC: saveStudent(Student) trên worker
        UC->>R: saveStudent(Student)
        R->>DB: insert/update
        DB-->>R: affected row
        R-->>VM: success
        VM-->>A: UiEvent + LiveData list
    end
""")
    add_mermaid(doc, "Hình 2.5 - Tuần tự ghi danh nhiều học viên", """
sequenceDiagram
    actor U as Người quản lý
    participant A as EnrollmentActivity
    participant VM as EnrollmentViewModel
    participant UC as EnrollmentUseCase
    participant R as Repository
    participant DB as SQLite
    U->>A: Chọn khóa và học viên
    A->>VM: enrollStudents(courseId, ids)
    VM->>UC: enroll(...) trên worker
    UC->>R: enrollStudents(ids, courseId)
    R->>DB: BEGIN + insert từng enrollment
    alt Tất cả thành công
        R->>DB: COMMIT
        VM-->>A: success event + summary mới
    else Có lỗi
        R->>DB: ROLLBACK
        VM-->>A: failure event
    end
""")

    start_page(doc, "2.8. Luồng thực hiện báo cáo", level=2)
    add_mermaid(doc, "Hình 2.6 - Tuần tự thực hiện báo cáo", """
sequenceDiagram
    actor U as Người quản lý
    participant A as ReportActivity
    participant VM as ReportViewModel
    participant UC as ReportUseCase
    participant R as ClassroomRepository
    participant DAO as ReportDao
    participant DB as SQLite
    U->>A: Chọn loại báo cáo và nhập bộ lọc
    A->>VM: runReport(input)
    VM->>UC: validate / prepare filter
    UC->>R: request report data
    R->>DAO: delegate query
    DAO->>DB: SELECT / JOIN / GROUP BY
    DB-->>DAO: Cursor
    DAO-->>VM: domain models
    VM-->>A: LiveData<List<ReportRow>>
    A-->>U: Danh sách / số lượng / empty state
""")
    add_body(doc, "Các truy vấn thành viên và tổng số theo khóa nằm trong EnrollmentDao. Truy vấn chưa ghi danh, khoảng tuổi, thống kê theo nhóm và lọc theo thời gian nằm trong ReportDao. Cách chia theo feature tránh tập trung toàn bộ SQL vào một lớp lớn.")

    # Trang thiết kế dữ liệu
    start_page(doc, "2.9. Sơ đồ thực thể quan hệ và thiết kế dữ liệu", level=2)
    add_mermaid(doc, "Hình 2.7 - Sơ đồ thực thể quan hệ", """
erDiagram
    STUDENTS ||--o{ ENROLLMENTS : registers
    COURSES ||--o{ ENROLLMENTS : contains
    STUDENTS {
        INTEGER id PK
        TEXT code UK
        TEXT name
        INTEGER age
        TEXT scratch_level
    }
    COURSES {
        INTEGER id PK
        TEXT code UK
        TEXT name
        TEXT language
        TEXT level
        TEXT start_date
        TEXT end_date
    }
    ENROLLMENTS {
        INTEGER student_id PK,FK
        INTEGER course_id PK,FK
        TEXT enrolled_at
    }
""")
    add_captioned_table(
        doc,
        "Bảng 2.2 - Ràng buộc dữ liệu",
        ["RÀNG BUỘC", "MỤC ĐÍCH"],
        [
            ("UNIQUE students.code / courses.code", "Ngăn trùng mã nghiệp vụ."),
            ("CHECK age BETWEEN 5 AND 18", "Bảo vệ miền dữ liệu tuổi."),
            ("CHECK language IN ('Scratch','Python')", "Giới hạn ngôn ngữ hợp lệ."),
            ("PRIMARY KEY(student_id, course_id)", "Ngăn ghi danh trùng."),
            ("FOREIGN KEY ... ON DELETE CASCADE", "Không để lại enrollment mồ côi."),
            ("index_enrollments_course_id", "Tăng tốc truy vấn học viên theo khóa."),
            ("index_enrollments_enrolled_at", "Tăng tốc lọc ghi danh theo thời gian."),
            ("CHECK start_date <= end_date", "Bảo vệ thứ tự thời gian của khóa học mới."),
        ],
        [3700, 5660],
    )
    add_body(doc, "Database version 5 quản lý migration: v3 bổ sung index theo khóa; v4 seed dữ liệu; v5 thêm ngày bắt đầu/kết thúc khóa, thời điểm ghi danh và index thời gian mà không xóa dữ liệu cũ.")

    # Trang 13
    start_page(doc, "CHƯƠNG 3. TRIỂN KHAI, KẾT QUẢ VÀ ĐÁNH GIÁ")
    doc.add_heading("3.1. Mô hình triển khai", level=2)
    add_mermaid(doc, "Hình 3.1 - Mô hình triển khai", """
flowchart TB
    DEV[Máy phát triển\nAndroid Studio + JDK 21]
    APK[APK debug/release]
    DEVICE[Thiết bị Android\nAndroid 7.0 trở lên]
    APP[Ứng dụng Lớp học lập trình nhí]
    DB[(LopHocLapTrinh.db\nSQLite local)]
    DEV -->|Gradle build| APK
    APK -->|ADB / cài APK| DEVICE
    DEVICE --> APP
    APP --> DB
""")
    add_body(doc, "Không có server, API hoặc dịch vụ nền bên ngoài. APK chứa UI, domain và data layer; database được SQLiteOpenHelper tạo trong vùng dữ liệu riêng của ứng dụng. Mô hình này đơn giản, hoạt động offline và không phát sinh chi phí vận hành máy chủ.")
    doc.add_heading("3.2. Yêu cầu môi trường", level=2)
    for text in [
        "Android Studio/Android SDK với compile SDK 36.1.",
        "JDK 21 cho Gradle; source compatibility Java 11.",
        "Thiết bị hoặc emulator API 24 trở lên.",
        "Gradle wrapper đi kèm dự án; không cần cài database riêng.",
    ]:
        add_bullet(doc, text, bullet_id)

    # Trang 14
    start_page(doc, "3.3. Các bước cài đặt và triển khai", level=2)
    for text in [
        "Mở thư mục dự án bằng Android Studio và xác nhận sdk.dir trong local.properties.",
        "Đồng bộ Gradle để tải Android Gradle Plugin và các thư viện AndroidX/Material.",
        "Chọn emulator hoặc kết nối thiết bị đã bật USB debugging.",
        "Chạy cấu hình app; Android Studio build và cài APK lên thiết bị.",
        "Ở lần chạy đầu, ClassroomApplication tạo repository; ClassroomDatabase tạo schema, index và dữ liệu mẫu.",
        "Kiểm tra bốn màn hình từ trang chủ: Học viên, Khóa học, Ghi danh, Báo cáo và truy vấn; nhập ngày theo yyyy-MM-dd.",
    ]:
        add_numbered(doc, text, number_id)
    doc.add_heading("Lệnh kiểm tra từ command line", level=2)
    add_mermaid(doc, "Khối lệnh 3.1 — Build và kiểm thử", """
# Windows PowerShell / Command Prompt
gradlew.bat compileDebugJavaWithJavac
gradlew.bat testDebugUnitTest
gradlew.bat lintDebug

# Tạo APK debug
gradlew.bat assembleDebug
""")
    add_callout(doc, "Lưu ý môi trường", "Nếu Gradle daemon giữ một JDK không hợp lệ, dừng daemon bằng gradlew.bat --stop rồi chạy lại với JAVA_HOME trỏ tới JDK đầy đủ.")

    # Trang 15
    start_page(doc, "3.4. Kết quả chức năng", level=2)
    add_captioned_table(
        doc,
        "Bảng 3.1 - Kết quả thực hiện chức năng",
        ["MÀN HÌNH", "KẾT QUẢ THỰC HIỆN", "TRẠNG THÁI"],
        [
            ("Trang chủ", "Điều hướng tới bốn nhóm chức năng.", "Đạt"),
            ("Học viên", "CRUD, validation, chọn/xóa nhiều, loading/empty state.", "Đạt"),
            ("Khóa học", "CRUD Scratch/Python, cấp độ, ngày bắt đầu/kết thúc, cascade ghi danh.", "Đạt"),
            ("Ghi danh", "Lọc ứng viên, tự lưu thời điểm, ghi danh/hủy nhiều bằng transaction.", "Đạt"),
            ("Báo cáo theo khóa", "Dropdown tìm khóa, danh sách và tổng số học viên.", "Đạt"),
            ("Tổng hợp khóa", "Tất cả khóa và số học viên, kể cả khóa trống.", "Đạt"),
            ("Tra cứu học viên", "Chưa ghi danh và khoảng tuổi tùy chọn 5-18.", "Đạt"),
            ("Thống kê", "Số khóa và lượt ghi danh theo ngôn ngữ/cấp độ.", "Đạt"),
            ("Theo thời gian", "Lọc ghi danh theo khoảng ngày và hiển thị thời điểm.", "Đạt"),
            ("Python cơ bản", "Lọc tuổi 10–12, Python, cấp độ Cơ bản, không trùng.", "Đạt"),
            ("MVVM/lifecycle", "State danh sách/selection giữ trong ViewModel; event consume một lần.", "Đạt"),
        ],
        [1800, 5960, 1600],
    )
    doc.add_heading("Luồng giao diện", level=2)
    for text in [
        "Trang chủ → chọn nhóm chức năng.",
        "Danh sách → thêm/sửa/xóa hoặc bật chế độ chọn nhiều.",
        "Ghi danh → chọn khóa → chọn học viên → xác nhận.",
        "Báo cáo → chọn loại → nhập bộ lọc tương ứng → nhấn Xem báo cáo.",
    ]:
        add_bullet(doc, text, bullet_id)
    add_callout(doc, "Minh chứng giao diện", "Có thể bổ sung ảnh chụp bốn màn hình vào bản nộp cuối nếu giảng viên yêu cầu minh chứng trực quan; nội dung chức năng trong báo cáo đã đối chiếu với source code hiện tại.")

    # Trang 16
    start_page(doc, "3.5. Kết quả kiểm thử và đánh giá", level=2)
    add_captioned_table(
        doc,
        "Bảng 3.2 - Kết quả kiểm thử",
        ["NHÓM KIỂM THỬ", "NỘI DUNG", "KẾT QUẢ"],
        [
            ("Unit test domain", "Validation tuổi/ngày; chuẩn bị Student/Course; lọc ứng viên; delegation use case và báo cáo.", "13 test - pass"),
            ("Build Java", "Biên dịch source main và unit test.", "BUILD SUCCESSFUL"),
            ("Android Lint", "Tĩnh hóa resource/API/i18n/dependency.", "0 error; warning không chặn build"),
            ("Instrumented test", "Transaction rollback, CourseSummary, truy vấn báo cáo và migration v5.", "Biên dịch đạt; cần device/emulator để chạy"),
        ],
        [2100, 5260, 2000],
    )
    doc.add_heading("Đánh giá kiến trúc", level=2)
    for text in [
        "ViewModel chỉ phụ thuộc use case; Activity không giữ repository.",
        "Domain không import Android/data/presentation; data implement domain interface.",
        "Callback database được chuyển về main thread trước khi cập nhật LiveData.",
        "State và one-shot event được phân biệt rõ bằng LiveData và UiEvent.",
        "Factory tập trung dependency wiring, tránh init(repository) trong View.",
    ]:
        add_bullet(doc, text, bullet_id)

    # Trang 17
    start_page(doc, "KẾT LUẬN, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN")
    doc.add_heading("Kết luận", level=2)
    add_body(doc, "Ứng dụng đã đáp ứng các yêu cầu quản lý học viên, khóa học, ghi danh và báo cáo trên Android. SQLite bảo đảm tính toàn vẹn dữ liệu; transaction xử lý an toàn các thao tác nhiều bản ghi. Kiến trúc MVVM/use-case/repository làm rõ trách nhiệm từng tầng và tạo nền tảng thuận lợi cho kiểm thử, bảo trì.")
    doc.add_heading("Hạn chế", level=2)
    for text in [
        "Dữ liệu chỉ nằm trên một thiết bị, chưa đồng bộ hoặc sao lưu đám mây.",
        "Chưa có đăng nhập, phân quyền và nhật ký thay đổi.",
        "Form/dialog nhiều bước chưa phục hồi đầy đủ sau process recreation.",
        "Repository vẫn trả boolean/long nên chưa mô tả chi tiết nguyên nhân lỗi.",
        "Chưa có bộ UI test tự động cho toàn bộ luồng CRUD.",
    ]:
        add_bullet(doc, text, bullet_id)
    doc.add_heading("Hướng phát triển", level=2)
    for text in [
        "Bổ sung SavedStateHandle và immutable UiState cho form phức tạp.",
        "Dùng result type chi tiết cho duplicate/not-found/storage error.",
        "Bổ sung Room, DI framework hoặc tách Gradle module khi quy mô tăng.",
        "Thêm xuất CSV/PDF, tìm kiếm nâng cao và đồng bộ API có xác thực.",
        "Xây dựng ViewModel test và Espresso test cho các hành trình quan trọng.",
    ]:
        add_bullet(doc, text, bullet_id)

    # Trang 18
    start_page(doc, "TÀI LIỆU THAM KHẢO")
    references = [
        "Android Developers. Guide to app architecture. https://developer.android.com/topic/architecture",
        "Android Developers. ViewModel overview. https://developer.android.com/topic/libraries/architecture/viewmodel",
        "Android Developers. LiveData overview. https://developer.android.com/topic/libraries/architecture/livedata",
        "Android Developers. Save data using SQLite. https://developer.android.com/training/data-storage/sqlite",
        "SQLite Documentation. Foreign Key Support. https://www.sqlite.org/foreignkeys.html",
        "SQLite Documentation. Transactions. https://www.sqlite.org/lang_transaction.html",
        "Oracle. Java Platform, Standard Edition Documentation. https://docs.oracle.com/en/java/",
        "Tài liệu nội bộ dự án: requirements.md; KIEN_TRUC_VA_LUONG_HOAT_DONG.md; CHUC_NANG_GIAO_DIEN.md.",
    ]
    for reference in references:
        add_numbered(doc, reference, number_id)
    doc.add_heading("Phụ lục: quy ước chuyển sơ đồ sang Visual Paradigm", level=2)
    add_body(doc, "Các khối Mermaid trong báo cáo là đặc tả logic. Khi vẽ lại bằng Visual Paradigm, giữ nguyên tên tác nhân, use case, lớp, participant, entity và quan hệ; có thể thay đổi vị trí để tối ưu bố cục nhưng không đổi chiều phụ thuộc hoặc cardinality.")
    add_callout(doc, "Thông tin cần hoàn thiện", "Thay [MÃ SINH VIÊN] và [HỌ VÀ TÊN SINH VIÊN] trên bìa trước khi nộp.")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
